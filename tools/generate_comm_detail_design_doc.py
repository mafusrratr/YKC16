from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/平台通信详细设计-中石化2.0_远洋V2_33007V2.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    set_cell_margins(table)


def set_font(run, name="SimSun", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True,
             color="2E74B5" if level <= 2 else "1F4D78")
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True)
        set_cell_shading(hdr[i], "F2F4F7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_font(r)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        set_table_width(table, widths)
    return table


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(sec, attr, Inches(1))
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "SimSun"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.10


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于C型控制器的充电机软件")
    set_font(r, size=22, bold=True, color="0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("平台通信详细设计")
    set_font(r, size=24, bold=True, color="0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("中石化 2.0 / YuanyangV2（远洋项目） / 33007_V2")
    set_font(r, size=16, bold=True, color="1F4D78")
    for _ in range(9):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("国电南瑞南京控制系统有限公司")
    set_font(r, size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年6月")
    set_font(r, size=12)
    doc.add_page_break()


def add_revision(doc):
    add_heading(doc, "修订记录", 1)
    add_table(doc, ["版本号", "修订内容", "修订人", "发布日期"], [
        ["V0.1", "根据当前仓库代码与平台协议需求整理中石化2.0、远洋项目、33007_V2 平台通信详细设计初稿", "", "2026-06-05"],
    ], [1.0, 3.7, 1.1, 1.2])
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "目 录", 1)
    for line in [
        "一、引言",
        "1.1 设计概述",
        "1.2 适用范围",
        "1.3 参考文档与代码依据",
        "二、软件设计概述",
        "2.1 平台通信模块定位",
        "2.2 总体架构与数据流",
        "2.3 内部 MQTT 主题契约",
        "三、功能详细设计",
        "3.1 平台接入-中石化2.0",
        "3.2 平台接入-YuanyangV2（远洋项目）",
        "3.3 平台接入-33007_V2",
        "四、配置、部署与测试要求",
    ]:
        add_para(doc, line)
    doc.add_page_break()


def section_intro(doc):
    add_heading(doc, "一、引言", 1)
    add_heading(doc, "1.1 设计概述", 2)
    add_para(doc, "本文档参考《详细设计-模版0528.docx》的章节组织方式，结合当前仓库中 `core/comm/Zhongshihua2.0`、`core/comm/YuanyangV2` 与 `core/comm/33007_V2` 的代码结构和平台协议需求，说明平台通信适配进程的详细设计。")
    add_para(doc, "平台通信进程统一命名为 `tcu_comm`，但按平台以独立目录、独立 Makefile 和独立协议实现隔离。平台侧帧号、字段、加密链路和状态机均封闭在对应平台目录内，内部业务仍通过统一 MQTT 主题与 `tcu_logic`、`pile_controller`、`meter`、`logger` 协作。")
    add_heading(doc, "1.2 适用范围", 2)
    add_bullets(doc, [
        "适用于中石化 2.0 平台通信模块：`core/comm/Zhongshihua2.0/`。",
        "适用于 YuanyangV2（远洋项目）平台通信模块：`core/comm/YuanyangV2/`。",
        "适用于基于当前多进程架构的联调、测试、维护和后续平台帧扩展。",
        "不覆盖主控 CAN2CCU、电表 DLT645 和 `tcu_logic` 状态机内部算法细节；这些模块仅作为通信适配的数据源和命令目标说明。",
    ])
    add_heading(doc, "1.3 参考文档与代码依据", 2)
    add_table(doc, ["序号", "资料/代码路径", "说明"], [
        ["1", "docs/TCU充电计费单元概要设计.md", "系统总体架构、进程职责、MQTT 主题契约。"],
        ["2", "docs/SOFTWARE_OVERVIEW_DESIGN.md", "平台适配层边界和多进程协作说明。"],
        ["3", "core/comm/Zhongshihua2.0/comm_process.h/.cpp", "中石化 2.0 通信模块源码。"],
        ["4", "core/comm/Zhongshihua2.0/doc/FRAME_TYPES_STATUS.md", "中石化 2.0 协议帧需求与处理范围。"],
        ["5", "core/comm/YuanyangV2/comm_process.h/.cpp", "远洋项目通信模块源码。"],
        ["6", "core/comm/YuanyangV2/doc/FRAME_TYPES_STATUS.md", "远洋项目首版平台接入需求范围。"],
        ["7", "core/comm/YuanyangV2/test/TEST_CASE_001_MAIN_FLOW.md", "远洋项目主链路验收用例。"],
        ["8", "core/comm/33007_V2/comm_process.h/.cpp", "33007_V2 通信模块源码。"],
        ["9", "core/comm/33007_V2/test/TEST_CASE_001_MAIN_FLOW.md", "33007_V2 主链路验收用例。"],
        ["10", "config/tcu_comm_zhongshihua2.ini、core/comm/YuanyangV2/doc/yuanyang_v2_comm_config.example.ini、core/comm/33007_V2/doc/33007_v2_comm_config.example.ini", "平台运行配置示例。"],
    ], [0.55, 3.5, 2.45])


def section_overview(doc):
    add_heading(doc, "二、软件设计概述", 1)
    add_heading(doc, "2.1 平台通信模块定位", 2)
    add_para(doc, "`tcu_comm` 是平台协议适配与路由进程，承担平台 TCP 会话、登录认证、心跳保活、协议帧编解码、平台命令到内部 MQTT 命令的转换、内部运行数据到平台上行帧的封装。")
    add_para(doc, "模块设计边界为：平台字段差异不进入 `tcu_logic`；`tcu_logic` 只消费统一的 `start_charge`、`stop_charge`、`record_cfm`、`power_ctrl` 等内部命令，并输出统一的状态、计费和交易事件。")
    add_heading(doc, "2.2 总体架构与数据流", 2)
    add_numbered(doc, [
        "平台下行报文经 TCP 接收缓存处理、帧完整性校验和必要的解密后，进入平台命令解析函数。",
        "解析成功的平台命令封装为 JSON，通过 `tcu/plat/{gun}/cmd` 或 `tcu/plat/{gun}/event` 发布。",
        "内部上行数据由 `tcu_comm` 订阅 `logic`、`pile`、`meter` 主题后缓存到每枪运行态结构。",
        "运行态缓存按平台要求周期上送实时数据，或在 `start_complete`、`stop_complete`、`update_record` 等事件到达时立即组帧上送。",
        "交易记录确认由平台应答触发，`tcu_comm` 发布 `record_cfm` 并调用 `LogSender::confirmTradeRecord(...)` 完成本地记录确认。",
    ])
    add_heading(doc, "2.3 内部 MQTT 主题契约", 2)
    add_table(doc, ["主题", "方向", "QoS", "生产者", "消费者", "说明"], [
        ["tcu/plat/{gun}/cmd", "平台到内部", "1", "tcu_comm", "tcu_logic", "平台远程启动、停止、功率调控、记录确认等命令。"],
        ["tcu/plat/{gun}/event", "平台到内部", "1", "tcu_comm", "tcu_logic/HMI", "平台在线状态、二维码/配置变更、setConfig 等事件。"],
        ["tcu/logic/{gun}/event", "内部到平台", "1", "tcu_logic", "tcu_comm", "状态变化、交易记录 update_record、业务事件。"],
        ["tcu/logic/{gun}/feeData", "内部到平台", "1", "tcu_logic", "tcu_comm", "计费过程、分时段电量/金额、费率模型 ID。"],
        ["tcu/pile/{gun}/data", "内部到平台", "0", "pile_controller", "tcu_comm", "遥信 yx、遥测 yc、桩侧实时状态。"],
        ["tcu/pile/{gun}/event", "内部到平台", "1", "pile_controller", "tcu_comm", "启动完成、停止完成、VIN 鉴权请求等流程事件。"],
        ["tcu/meter/{gun}/data", "内部到平台", "0", "meter", "tcu_comm", "电表累计电量、电压、电流。"],
    ], [1.55, 0.7, 0.45, 0.9, 0.9, 2.0])


def add_fault_record_code_table(doc, comm_prefix, comm_rows):
    rows = [
        ["pile_controller", "PILE_FAULT_001", "桩侧遥信故障", "pile `yx.otherFault` 位图存在有效故障位，或 YX22/YX23 聚合后判定为设备故障。", "由 `tcu_logic` 归一化为常态故障，发布 `tcu/save/{gun}/event type=Error`；平台实时帧同步故障状态。"],
        ["pile_controller", "PILE_FAULT_002", "启动完成失败", "`tcu/pile/{gun}/event` 中 `start_complete.successFlag != 0`，或携带 `chargeFailReason`。", "按启动失败类故障生成统一 `reason`，写入本次充电记录，不重复发布常态 Error。"],
        ["pile_controller", "PILE_FAULT_003", "停止完成异常", "`stop_complete.stopReason`、`bmsStopReason`、`bmsChargeFaultReason`、`bmsStopErrorReason` 命中故障停机规则。", "按充电中失败类故障生成统一 `reason`，写入交易记录和结算展示。"],
        ["meter", "METER_FAULT_001", "电表通信中断", "电表轮询超时、连续读取失败或 meter 在线状态转离线。", "作为常态故障或充电中计量异常上报；空闲/准备阶段发布 Error，交易中写入统一 `reason` 或计量异常标记。"],
        ["meter", "METER_FAULT_002", "电表数据异常", "累计电量回退、读数跳变、电压/电流无效或超出合理范围。", "由 `tcu_logic` 判定是否影响计费；影响交易时写入记录异常原因，并保留原始读数用于追溯。"],
        ["tcu_logic", "LOGIC_FAULT_001", "启动状态超时", "进入 STARTING 后超时未收到 `start_complete`。", "按启动失败故障生成统一 `reason=REASON_START_TIMEOUT`，停止启动流程并生成记录。"],
        ["tcu_logic", "LOGIC_FAULT_002", "停止状态超时", "进入 STOPPING 后超时未收到 `stop_complete` 或末端计量未收敛。", "按控制类异常记录，必要时触发补停机和记录异常标记。"],
        ["tcu_logic", "LOGIC_FAULT_003", "平台/本地命令状态不匹配", "收到启动、停止、记录确认等命令时当前状态不允许执行。", "拒绝命令或保持原状态，记录业务日志，避免重复启动/重复停机。"],
        ["logger", "LOGGER_FAULT_001", "交易记录落库失败", "SQLite 写入失败、数据库不可用或字段不完整。", "保留日志告警，交易记录保持待确认/待补写状态，避免平台确认提前闭环。"],
        ["logger", "LOGGER_FAULT_002", "交易记录确认失败", "平台确认后本地 `confirmTradeRecord` 无法匹配或更新失败。", "保留未确认状态，后续补确认或人工排查。"],
    ]
    rows.extend(comm_rows)
    add_table(doc, ["模块/来源", "代码", "故障/异常名称", "触发条件", "处理要求"], rows, [1.0, 1.25, 1.35, 2.05, 2.4])


def section_zhong(doc):
    add_heading(doc, "三、功能详细设计", 1)
    add_heading(doc, "3.1 平台接入-中石化2.0", 2)
    add_para(doc, "中石化 2.0 平台接入模块负责完成平台 TCP 通信、登录鉴权、费率同步、远程启停、并充/VIN 启动、运行数据上报、交易记录上送与确认等功能。平台协议需求封装在 `core/comm/Zhongshihua2.0` 内部，对内统一使用 `tcu/plat/{gun}/cmd|event` 与业务逻辑交互。")

    add_heading(doc, "3.1.1 模块架构", 3)
    add_heading(doc, "3.1.1.1 关键结构体设计", 3)
    add_table(doc, ["序号", "结构体名称", "所属头文件", "职责"], [
        ["1", "CommConfig", "comm_process.h", "保存枪数、MQTT、平台 TCP、桩编号、登录 ID、MAC、企业信用代码、SM2 公钥、二维码、枪 ID、调试开关等运行配置。"],
        ["2", "StartCompleteData", "comm_process.h", "缓存启动完成及 BMS 基础数据，供启动结果、BRM、BCP 等平台上行帧取数。"],
        ["3", "FeeSegmentData", "comm_process.h", "保存分时段电量、电费、服务费，供实时数据和交易记录上报。"],
        ["4", "GunRuntimeData", "comm_process.h", "按枪汇总启动参数、遥信、遥测、遥脉、电表、计费模型、VIN 鉴权和交易记录状态。"],
        ["5", "PlatformLoginState", "comm_process.h", "描述平台登录、计费模型同步、对时、在线心跳等链路状态。"],
        ["6", "FeeModel", "base/logger/log_sender.h", "统一描述平台费率模型，包括时段、尖峰平谷标识、电价与服务费。"],
    ], [0.55, 1.4, 1.55, 3.0])
    add_heading(doc, "3.1.1.2 关键子功能划分", 3)
    add_table(doc, ["序号", "子功能", "实现文件", "核心函数/接口"], [
        ["1", "配置加载与运行态初始化", "comm_process.cpp", "loadConfig、doInitialize"],
        ["2", "内部 MQTT 订阅与路由", "comm_process.cpp", "initMqtt、onMqttConnected、onMqttMessage、parseTopic、publishPlatCommand、publishSetConfig"],
        ["3", "平台 TCP 连接维护", "comm_process.cpp", "connectPlatformTcp、maintainPlatformTcp、closePlatformTcp、handlePlatformRxData"],
        ["4", "二进制帧封装与校验", "comm_process.cpp", "buildPlatformFrame、sendPlatformFrame、processPlatformPacket、calcCrc16Modbus"],
        ["5", "登录与安全上下文", "comm_process.cpp", "driveLoginStateMachine、prepareLoginCryptoContext、tryUpdateSm2PubKeyFromLoginAck"],
        ["6", "平台命令解析", "comm_process.cpp", "parseRemoteStart0A8、parseRemoteMergeStart0A4、parseRemoteStop036、parseRecordConfirm040、parseQrCodeSet05A"],
        ["7", "业务上行组帧", "comm_process.cpp", "buildChargeInfoBody、buildStartChargeResultBody、buildBrmBody、buildBcpBody、buildChargeRecordBodyFromUpdateRecord"],
        ["8", "计费模型与记录确认", "comm_process.cpp", "parseFeeModelAck00A、loadFeeModelFromDbFile、refreshTradeRecordFeeModelCache、LogSender::saveFeeModel、confirmTradeRecord"],
    ], [0.55, 1.55, 1.55, 3.0])

    add_heading(doc, "3.1.2 状态机实现", 3)
    add_heading(doc, "3.1.2.1 状态枚举定义", 3)
    add_table(doc, ["状态枚举", "状态名称", "需求说明"], [
        ["LOGIN_IDLE", "空闲/待认证", "TCP 已连接或重连完成后进入该状态，准备登录安全上下文并发起平台认证。"],
        ["LOGIN_REQ_AUTH", "登录认证中", "周期发送 0x01 登录认证请求，等待 0x02 登录认证应答。"],
        ["LOGIN_REQ_FEE_MODEL", "计费模型同步中", "认证成功后按枪发送计费模型请求，等待平台返回有效费率。"],
        ["LOGIN_REQ_TIME_SYNC", "对时中", "计费模型同步后发送对时请求，完成平台时间同步。"],
        ["LOGIN_ONLINE", "在线运行", "维持心跳，接收平台业务命令，并按需求上报运行数据和交易记录。"],
    ], [1.45, 1.3, 3.35])
    add_heading(doc, "3.1.2.2 状态转换矩阵", 3)
    add_table(doc, ["当前状态", "触发事件", "目标状态", "处理要求"], [
        ["LOGIN_IDLE", "TCP 建链成功且允许发起登录", "LOGIN_REQ_AUTH", "准备 SM2/SM4 相关上下文，构造 0x01 登录认证帧。"],
        ["LOGIN_REQ_AUTH", "收到 0x02 且认证成功", "LOGIN_REQ_FEE_MODEL", "解析认证结果，必要时更新平台公钥，进入费率同步。"],
        ["LOGIN_REQ_AUTH", "认证失败或帧解析失败", "LOGIN_IDLE", "记录失败原因，按重试间隔限流后重新登录。"],
        ["LOGIN_REQ_FEE_MODEL", "收到有效 0x0A 费率模型", "LOGIN_REQ_TIME_SYNC", "保存 FeeModel，更新按枪费率缓存。"],
        ["LOGIN_REQ_TIME_SYNC", "收到 0x0C 对时应答", "LOGIN_ONLINE", "执行系统对时并发布平台在线事件。"],
        ["LOGIN_ONLINE", "心跳超时、发送失败或 TCP 关闭", "LOGIN_IDLE", "关闭 socket，发布离线事件，等待重连。"],
    ], [1.35, 1.7, 1.35, 2.3])

    add_heading(doc, "3.1.3 订阅主题列表", 3)
    add_table(doc, ["序号", "主题", "QoS", "用途"], [
        ["1", "tcu/logic/+/event", "1", "接收状态变化、VIN 启动申请、交易记录 update_record 等业务事件。"],
        ["2", "tcu/logic/+/feeData", "1", "接收分时段计费、电量、金额、费率模型 ID 等计费数据。"],
        ["3", "tcu/pile/+/data", "0", "接收遥信 yx、遥测 yc、枪状态、BMS 请求与实测数据。"],
        ["4", "tcu/pile/+/event", "1", "接收 start_complete、stop_complete、VIN 识别等流程事件。"],
        ["5", "tcu/meter/+/data", "0", "接收电表示数、电压、电流等计量数据。"],
    ], [0.55, 2.0, 0.55, 3.4])

    add_heading(doc, "3.1.4 关键数据结构详细定义", 3)
    add_table(doc, ["结构体", "字段", "类型/单位", "详细说明"], [
        ["CommConfig", "gunCount、masterHost、masterPort、cdzNo、loginId、macAddr、factoryCreditCode、sm2PublicKey、gunQrCodeList、gunIdList", "配置项", "描述平台通信、桩身份、安全材料、枪号和二维码等静态参数。"],
        ["StartCompleteData", "successFlag、failReason、batteryType、ratedCapacity、soc、vin、bmsSoftwareVersion", "业务缓存", "保存启动完成后平台所需的车辆/BMS 数据。"],
        ["FeeSegmentData", "startTs、endTs、energyKwh、electricAmount、serviceAmount", "kWh/元", "保存每个计费时段的电量和费用明细。"],
        ["GunRuntimeData", "gunStatus、yxWorkStatus、voltage、current、soc、meterEnergy、totalAmount、feeModelId、pendingVinAuth*", "混合运行态", "按枪聚合内部多主题数据，为平台实时帧和交易帧提供统一数据源。"],
        ["FeeModel", "feeModelId、timeNum、timeSeg、segFlag、chargeFee、serviceFee", "费率模型", "平台下发或启动携带的费率数据，落库后供计费和记录上送引用。"],
    ], [1.25, 2.35, 1.0, 2.4])

    add_heading(doc, "3.1.5 接口详细定义", 3)
    add_table(doc, ["接口/函数", "输入", "输出", "接口说明"], [
        ["buildPlatformFrame", "cmd、body、seqOverride", "二进制平台帧", "统一封装中石化 2.0 TCP 帧头、时间、长度、序号、命令字、信息体和 CRC。"],
        ["processPlatformPacket", "完整平台帧", "内部命令或状态动作", "完成命令字识别、解密/校验和业务分发。"],
        ["publishPlatCommand", "gun、cmd、dataObj", "MQTT JSON", "向 `tcu/plat/{gun}/cmd` 发布统一平台命令。"],
        ["publishSetConfig", "gun、dataObj", "MQTT JSON retain", "向 `tcu/plat/{gun}/event` 发布二维码、平台链路等配置事件。"],
        ["parseRemoteStart0A8", "平台 0xA8 信息体", "gun、start_charge data、FeeModel", "解析远程启动请求和随包费率要求。"],
        ["parseRemoteStop036", "平台 0x36 信息体", "gun、stop_charge data", "解析平台远程停机请求。"],
        ["buildChargeRecordBodyFromUpdateRecord", "gun、update_record data", "0x3D 信息体", "将逻辑侧交易记录转换为平台交易记录帧。"],
    ], [1.6, 1.55, 1.55, 2.1])

    add_heading(doc, "3.1.6 故障记录代码清单", 3)
    add_fault_record_code_table(doc, "ZSH", [
        ["comm/Zhongshihua2.0", "ZSH_COMM_001", "平台 TCP 连接失败", "主站不可达、connect 失败或 socket 异常关闭。", "记录日志，按 `tcp_reconnect_sec` 重连，并向内部发布平台离线事件。"],
        ["comm/Zhongshihua2.0", "ZSH_COMM_002", "登录认证失败", "0x02 应答失败、认证帧解析失败或安全上下文异常。", "回到 LOGIN_IDLE，按 `loginRetrySec` 重试。"],
        ["comm/Zhongshihua2.0", "ZSH_COMM_003", "计费模型解析失败", "0x0A 或 0x58 费率字段不完整、时段非法或价格解析失败。", "拒绝或不生效该费率，保留原费率缓存。"],
        ["comm/Zhongshihua2.0", "ZSH_COMM_004", "平台帧校验失败", "帧头、长度、CRC、枪号或桩号校验失败。", "丢弃该帧，必要时回失败应答。"],
        ["comm/Zhongshihua2.0", "ZSH_COMM_005", "心跳超时", "在线态长时间未收到 0x04 或发送心跳失败。", "关闭 TCP，发布平台离线事件。"],
        ["comm/Zhongshihua2.0", "ZSH_COMM_006", "交易记录确认失败", "0x40 返回失败或确认记录号无法匹配。", "保留本地未确认状态，等待后续补传或人工排查。"],
    ])


def section_yuanyang(doc):
    add_heading(doc, "3.2 平台接入-YuanyangV2（远洋项目）", 2)
    add_para(doc, "YuanyangV2 平台接入模块负责完成远洋运营前置通信规约要求，包括 RSA/AES 握手、心跳、总召唤、三遥与遥脉上送、故障上送、启停控制、功率调控、计费模型 2.0、交易记录上送与确认。模块对内仍保持统一 MQTT 主题契约。")

    add_heading(doc, "3.2.1 模块架构", 3)
    add_heading(doc, "3.2.1.1 关键结构体设计", 3)
    add_table(doc, ["序号", "结构体名称", "所属头文件", "职责"], [
        ["1", "CommConfig", "comm_process.h", "保存 MQTT、平台 TCP、站地址、设备地址、资产编码、RSA 公钥、离线运行和调试配置。"],
        ["2", "GunRuntimeData", "comm_process.h", "按枪缓存遥信、遥测、电表、计费、订单和交易记录状态。"],
        ["3", "ParsedFrame", "comm_process.h", "描述平台帧解析结果，包括控制域、源/目标地址、ASDU、VSQ、COT、公共地址和尾部数据。"],
        ["4", "PendingControl", "comm_process.h", "保存平台启停控制待确认状态，等待桩侧完成事件后回确认或拒绝。"],
        ["5", "PlatformLoginState", "comm_process.h", "描述 RSA 公钥请求、设备认证、在线运行等平台链路状态。"],
        ["6", "FeeSegmentData/FeeModel", "comm_process.h / log_sender.h", "描述分时段计费明细及平台计费模型 2.0。"],
    ], [0.55, 1.45, 1.6, 2.95])
    add_heading(doc, "3.2.1.2 关键子功能划分", 3)
    add_table(doc, ["序号", "子功能", "实现文件", "核心函数/接口"], [
        ["1", "配置加载与地址解析", "comm_process.cpp", "loadConfig、parseAssetAddressCode、gunFromDeviceAddr、deviceAddrFromGun"],
        ["2", "MQTT 订阅与内部路由", "comm_process.cpp", "initMqtt、onMqttConnected、onMqttMessage、publishPlatCommand、publishPlatformLinkEvent"],
        ["3", "TCP 与心跳维护", "comm_process.cpp", "connectPlatformTcp、maintainPlatformTcp、buildHeartbeatFrame、sendHeartbeat"],
        ["4", "ASDU 帧封装解析", "comm_process.cpp", "buildAsduFrame、parseFrame、processPlatformPacket、sendAsdu"],
        ["5", "RSA/AES 安全链路", "comm_process.cpp", "sendRsaPublicKeyRequest、sendDeviceAuthRequest、encryptTail、decryptTail、tryUpdateRsaPubKeyFromResponse"],
        ["6", "平台下行解析", "comm_process.cpp", "parseRemoteControl041、parsePowerControl02E、parseFeeModel051、parseRecordConfirm042"],
        ["7", "三遥遥脉与故障上送", "comm_process.cpp", "buildAllTelesignalTail、buildAllTelemetryTail、buildLegacyPulseTail、buildExtPulseTail、buildFaultTail、reportRuntimePeriodic"],
        ["8", "交易记录上送", "comm_process.cpp", "buildChargeRecordTail、LogSender::confirmTradeRecord"],
    ], [0.55, 1.55, 1.55, 3.0])

    add_heading(doc, "3.2.2 状态机实现", 3)
    add_heading(doc, "3.2.2.1 状态枚举定义", 3)
    add_table(doc, ["状态枚举", "状态名称", "需求说明"], [
        ["LOGIN_IDLE", "空闲/待握手", "TCP 已连接后准备会话密钥，等待发起 RSA 公钥请求。"],
        ["LOGIN_REQ_RSA_PUBLIC_KEY", "公钥请求中", "发送 0x81 请求平台 RSA 公钥，等待平台返回可用公钥。"],
        ["LOGIN_REQ_AUTH", "设备认证中", "使用 RSA 加密 AES 会话密钥并发送 0x82 认证请求。"],
        ["LOGIN_ONLINE", "在线运行", "业务帧按 AES 处理，维持心跳，处理总召、启停、功率、费率、记录确认等业务。"],
    ], [1.55, 1.25, 3.35])
    add_heading(doc, "3.2.2.2 状态转换矩阵", 3)
    add_table(doc, ["当前状态", "触发事件", "目标状态", "处理要求"], [
        ["LOGIN_IDLE", "TCP 建链成功", "LOGIN_REQ_RSA_PUBLIC_KEY", "生成 AES 会话密钥，发送 0x81 公钥请求。"],
        ["LOGIN_REQ_RSA_PUBLIC_KEY", "收到有效 RSA 公钥", "LOGIN_REQ_AUTH", "更新并固化公钥，发送设备认证请求。"],
        ["LOGIN_REQ_RSA_PUBLIC_KEY", "公钥解析失败", "LOGIN_IDLE", "记录异常，按重试间隔重新请求。"],
        ["LOGIN_REQ_AUTH", "收到 0x82 COT 确认", "LOGIN_ONLINE", "AES 会话密钥生效，发布在线事件并请求 0x51 费率。"],
        ["LOGIN_REQ_AUTH", "认证拒绝或解密失败", "LOGIN_IDLE", "清空会话状态，重新握手。"],
        ["LOGIN_ONLINE", "收到 0x83 重置密钥", "LOGIN_REQ_RSA_PUBLIC_KEY", "清空 AES 会话，重新获取公钥和认证。"],
        ["LOGIN_ONLINE", "心跳超时或 TCP 断开", "LOGIN_IDLE", "关闭 socket，发布离线事件，等待重连。"],
    ], [1.3, 1.65, 1.45, 2.35])

    add_heading(doc, "3.2.3 订阅主题列表", 3)
    add_table(doc, ["序号", "主题", "QoS", "用途"], [
        ["1", "tcu/logic/+/event", "1", "接收状态变化、交易记录 update_record，以及启停结果相关业务事件。"],
        ["2", "tcu/logic/+/feeData", "1", "接收总电量、总金额、分时段电量/金额和费率模型信息。"],
        ["3", "tcu/pile/+/data", "0", "接收遥信、遥测、枪状态、故障、BMS 需求和实测数据。"],
        ["4", "tcu/pile/+/event", "1", "接收 start_complete、stop_complete，用于 0x41 控制确认。"],
        ["5", "tcu/meter/+/data", "0", "接收电表示数、电压、电流，用于遥测和遥脉上送。"],
    ], [0.55, 2.0, 0.55, 3.4])

    add_heading(doc, "3.2.4 关键数据结构详细定义", 3)
    add_table(doc, ["结构体", "字段", "类型/单位", "详细说明"], [
        ["CommConfig", "gunCount、masterHost、masterPort、stationAddr、assetCode、deviceAddrList、rsaPublicKey", "配置项", "描述平台地址、安全公钥、站设备地址和内部 MQTT 参数。"],
        ["ParsedFrame", "control、targetStation、targetDevice、sourceStation、sourceDevice、cmd、vsq、cot、commonAddr、tail", "协议解析结构", "保存类 104 平台帧解析结果，供业务分发使用。"],
        ["PendingControl", "active、action、vsq、tailPlain、createdAt", "待确认控制", "保存 0x41 启停命令上下文，桩侧完成后按原 VSQ 和尾部回确认。"],
        ["GunRuntimeData", "gunStatus、yx*、voltage、current、soc、meterEnergy、totalAmount、feeSegments、pendingRecordTradeNo", "混合运行态", "作为三遥、遥脉、故障、交易记录上送的数据源。"],
        ["FeeModel", "feeModelId、timeNum、timeSeg、segFlag、chargeFee、serviceFee", "费率模型", "由 0x51 计费模型 2.0 下发解析生成，保存后用于计费和遥脉统计。"],
    ], [1.25, 2.45, 1.0, 2.3])

    add_heading(doc, "3.2.5 接口详细定义", 3)
    add_table(doc, ["接口/函数", "输入", "输出", "接口说明"], [
        ["buildAsduFrame", "cmd、vsq、cot、tailPlain、sourceAddr", "平台 ASDU 帧", "封装远洋平台业务帧，按要求补充控制域、地址域、ASDU 和加密尾部。"],
        ["parseFrame", "完整平台帧", "ParsedFrame", "解析控制域、站设备地址、命令字、VSQ、COT 和尾部数据。"],
        ["encryptTail/decryptTail", "cmd、明文/密文 tail", "密文/明文 tail", "按握手状态执行 RSA/AES 加解密处理。"],
        ["parseRemoteControl041", "0x41 ParsedFrame", "gun、action、dataObj", "解析平台启停控制，生成内部 start_charge 或 stop_charge 数据。"],
        ["parsePowerControl02E", "0x2E ParsedFrame", "gun、dataObj", "解析功率调控命令并转内部 power_ctrl。"],
        ["parseFeeModel051", "0x51 ParsedFrame", "FeeModel", "解析平台计费模型 2.0，并转换为内部费率结构。"],
        ["buildChargeRecordTail", "gun、update_record data", "0x42 tail", "将逻辑侧交易记录转换为远洋平台交易记录。"],
        ["sendControlAck", "gun、PendingControl、COT", "0x41 确认帧", "根据桩侧完成结果向平台回确认或拒绝。"],
    ], [1.55, 1.55, 1.55, 2.15])

    add_heading(doc, "3.2.6 故障记录代码清单", 3)
    add_fault_record_code_table(doc, "YY", [
        ["comm/YuanyangV2", "YY_COMM_001", "平台 TCP 连接失败", "主站不可达、connect 失败或 socket 异常关闭。", "记录日志，按 `tcp_reconnect_sec` 重连，并向内部发布平台离线事件。"],
        ["comm/YuanyangV2", "YY_COMM_002", "RSA 公钥获取失败", "0x81 应答缺失、公钥格式非法或固化失败。", "回到 LOGIN_IDLE，重新请求平台公钥。"],
        ["comm/YuanyangV2", "YY_COMM_003", "设备认证失败", "0x82 被拒绝、RSA 加密失败或 AES 会话异常。", "清空会话密钥，重新握手。"],
        ["comm/YuanyangV2", "YY_COMM_004", "平台帧解析/解密失败", "长度、地址、COT、AES 解密或 tail 解析失败。", "丢弃该帧，必要时回 COT 0x06 拒绝。"],
        ["comm/YuanyangV2", "YY_COMM_005", "控制确认超时", "0x41 启停命令已下发但超时未收到 start_complete/stop_complete。", "回 COT 0x06，清理 PendingControl。"],
        ["comm/YuanyangV2", "YY_COMM_006", "计费模型解析失败", "0x51 费率时段、尖峰平谷标识或价格字段非法。", "回拒绝确认，保持原费率模型。"],
        ["comm/YuanyangV2", "YY_COMM_007", "交易记录确认失败", "0x42 确认帧无法匹配本地 pendingRecordTradeNo。", "保留未确认记录，等待后续补传或人工排查。"],
    ])


def section_33007(doc):
    add_heading(doc, "3.3 平台接入-33007_V2", 2)
    add_para(doc, "33007_V2 平台接入模块面向 NBT33007-2013 执行细则的 V2 版本通信需求，负责完成平台 TCP 建链、控制域心跳、总召唤、普通/扩展三遥、普通/扩展遥脉、故障信息、启停控制、功率调控、计费模型、交易记录上送与确认。该版本业务 ASDU 按明文处理，不使用远洋项目中的 RSA/AES 安全握手。")

    add_heading(doc, "3.3.1 模块架构", 3)
    add_heading(doc, "3.3.1.1 关键结构体设计", 3)
    add_table(doc, ["序号", "结构体名称", "所属头文件", "职责"], [
        ["1", "CommConfig", "comm_process.h", "保存枪数、MQTT、平台 TCP、站地址、资产编码、设备地址列表、离线运行和调试配置。"],
        ["2", "GunRuntimeData", "comm_process.h", "按枪缓存遥信、遥测、电表、计费、订单和交易记录待确认状态。"],
        ["3", "ParsedFrame", "comm_process.h", "描述平台帧解析结果，包括控制域、源/目标地址、ASDU、VSQ、COT、公共地址和明文尾部数据。"],
        ["4", "PendingControl", "comm_process.h", "保存平台 0x41 启停控制待确认状态，等待桩侧完成事件后回确认或拒绝。"],
        ["5", "PlatformLoginState", "comm_process.h", "描述平台 TCP 在线状态；RSA/AES 相关枚举保留兼容，但 33007_V2 主链路建连后直接进入在线。"],
        ["6", "FeeSegmentData/FeeModel", "comm_process.h / log_sender.h", "描述分时段计费明细及 0x50 计费模型。"],
    ], [0.55, 1.45, 1.6, 2.95])

    add_heading(doc, "3.3.1.2 关键子功能划分", 3)
    add_table(doc, ["序号", "子功能", "实现文件", "核心函数/接口"], [
        ["1", "配置加载与站设备地址解析", "comm_process.cpp", "loadConfig、parseAssetAddressCode、gunFromDeviceAddr、deviceAddrFromGun"],
        ["2", "MQTT 订阅与内部路由", "comm_process.cpp", "initMqtt、onMqttConnected、onMqttMessage、publishPlatCommand、publishPlatformLinkEvent"],
        ["3", "TCP 建链与心跳维护", "comm_process.cpp", "connectPlatformTcp、connectWithTimeout、maintainPlatformTcp、buildHeartbeatFrame、sendHeartbeat"],
        ["4", "明文 ASDU 帧封装解析", "comm_process.cpp", "buildAsduFrame、parseFrame、processPlatformPacket、sendAsdu"],
        ["5", "平台下行命令解析", "comm_process.cpp", "parseRemoteControl041、parsePowerControl02E、parseFeeModel050、parseRecordConfirm042"],
        ["6", "普通/扩展三遥上送", "comm_process.cpp", "buildAllLegacyTelesignalTail、buildAllTelesignalTail、buildAllLegacyTelemetryTail、buildAllTelemetryTail、reportRuntimePeriodic"],
        ["7", "普通/扩展遥脉与故障上送", "comm_process.cpp", "buildLegacyPulseTail、buildExtPulseTail、buildFaultTail、sendTotalCallResponse"],
        ["8", "交易记录上送与确认", "comm_process.cpp", "buildChargeRecordTail、LogSender::confirmTradeRecord"],
    ], [0.55, 1.55, 1.55, 3.0])

    add_heading(doc, "3.3.2 状态机实现", 3)
    add_heading(doc, "3.3.2.1 状态枚举定义", 3)
    add_table(doc, ["状态枚举", "状态名称", "需求说明"], [
        ["LOGIN_IDLE", "离线/待连接", "平台 TCP 未连接或已关闭，等待重连。"],
        ["LOGIN_ONLINE", "在线运行", "TCP 建链成功后进入在线状态，发布平台在线事件，发送心跳并请求 0x50 计费模型。"],
        ["LOGIN_REQ_RSA_PUBLIC_KEY", "兼容保留", "33007_V2 标准主链路不使用 RSA 公钥请求，仅为复用结构保留。"],
        ["LOGIN_REQ_AUTH", "兼容保留", "33007_V2 标准主链路不使用设备加密认证，仅为复用结构保留。"],
    ], [1.55, 1.25, 3.35])

    add_heading(doc, "3.3.2.2 状态转换矩阵", 3)
    add_table(doc, ["当前状态", "触发事件", "目标状态", "处理要求"], [
        ["LOGIN_IDLE", "TCP 建链成功", "LOGIN_ONLINE", "发布平台在线事件，初始化心跳时间，主动发送 0x50 计费模型请求。"],
        ["LOGIN_IDLE", "TCP 建链失败", "LOGIN_IDLE", "按 `tcp_reconnect_sec` 重试，并持续喂守护进程看门狗。"],
        ["LOGIN_ONLINE", "达到心跳周期", "LOGIN_ONLINE", "发送控制域 0x43 心跳帧，等待平台 0x83 心跳应答。"],
        ["LOGIN_ONLINE", "收到 0x03 总召唤", "LOGIN_ONLINE", "按需求回复普通遥测、扩展遥测、普通遥信、扩展遥信、普通遥脉、扩展遥脉和故障信息。"],
        ["LOGIN_ONLINE", "心跳应答超时或 TCP 关闭", "LOGIN_IDLE", "关闭 socket，发布平台离线事件，等待重连。"],
    ], [1.3, 1.65, 1.45, 2.35])

    add_heading(doc, "3.3.3 订阅主题列表", 3)
    add_table(doc, ["序号", "主题", "QoS", "用途"], [
        ["1", "tcu/logic/+/event", "1", "接收状态变化、交易记录 update_record，以及启停结果相关业务事件。"],
        ["2", "tcu/logic/+/feeData", "1", "接收总电量、总金额、分时段电量/金额和费率模型信息。"],
        ["3", "tcu/pile/+/data", "0", "接收遥信、遥测、枪状态、故障、BMS 需求和实测数据。"],
        ["4", "tcu/pile/+/event", "1", "接收 start_complete、stop_complete，用于 0x41 控制确认。"],
        ["5", "tcu/meter/+/data", "0", "接收电表示数、电压、电流，用于遥测和遥脉上送。"],
    ], [0.55, 2.0, 0.55, 3.4])

    add_heading(doc, "3.3.4 关键数据结构详细定义", 3)
    add_table(doc, ["结构体", "字段", "类型/单位", "详细说明"], [
        ["CommConfig", "gunCount、masterHost、masterPort、stationAddr、assetCode、deviceAddrList、offlineRunMode、debugTcp", "配置项", "描述平台地址、站设备地址、内部 MQTT 参数和调试运行模式。"],
        ["ParsedFrame", "control、targetStation、targetDevice、sourceStation、sourceDevice、cmd、vsq、cot、commonAddr、tail", "协议解析结构", "保存明文 ASDU 平台帧解析结果，供业务分发使用。"],
        ["PendingControl", "active、action、vsq、tailPlain、createdAt", "待确认控制", "保存 0x41 启停命令上下文，按原命令体回 COT 0x04/0x06。"],
        ["GunRuntimeData", "gunStatus、yx*、voltage、current、soc、meterEnergy、totalAmount、feeSegments、pendingRecordTradeNo", "混合运行态", "作为普通/扩展三遥、遥脉、故障、交易记录上送的数据源。"],
        ["FeeModel", "feeModelId、timeNum、timeSeg、segFlag、chargeFee、serviceFee", "费率模型", "由 0x50 计费模型下发解析生成；远程启动时可附加到 start_charge 内部命令。"],
    ], [1.25, 2.45, 1.0, 2.3])

    add_heading(doc, "3.3.5 接口详细定义", 3)
    add_table(doc, ["接口/函数", "输入", "输出", "接口说明"], [
        ["connectWithTimeout", "socket、地址、超时时间", "连接结果", "限制 TCP 建链阻塞时间，避免守护进程看门狗饥饿。"],
        ["buildAsduFrame", "cmd、vsq、cot、tailPlain、sourceAddr", "平台 ASDU 帧", "封装 33007_V2 明文业务帧，补充控制域、源/目标地址、ASDU 和长度。"],
        ["parseFrame", "完整平台帧", "ParsedFrame", "解析控制域、站设备地址、命令字、VSQ、COT 和明文尾部。"],
        ["parseRemoteControl041", "0x41 ParsedFrame", "gun、action、dataObj", "解析平台启停控制，生成内部 start_charge 或 stop_charge 数据；启动时要求具备可用费率模型。"],
        ["parsePowerControl02E", "0x2E ParsedFrame", "gun、dataObj", "解析功率调控命令并转内部 power_ctrl。"],
        ["parseFeeModel050", "0x50 ParsedFrame", "FeeModel", "解析平台计费模型，并转换为内部费率结构。"],
        ["sendTotalCallResponse", "总召唤事件", "多类上行 ASDU", "按需求响应普通遥测、扩展遥测、普通遥信、扩展遥信、遥脉和故障。"],
        ["buildChargeRecordTail", "gun、update_record data", "0x42 tail", "将逻辑侧交易记录转换为 33007_V2 平台交易记录。"],
    ], [1.55, 1.55, 1.55, 2.15])

    add_heading(doc, "3.3.6 故障记录代码清单", 3)
    add_fault_record_code_table(doc, "33007", [
        ["comm/33007_V2", "33007_COMM_001", "平台 TCP 连接失败", "主站不可达、connect 超时或 socket 异常关闭。", "记录日志，按 `tcp_reconnect_sec` 重连，并向内部发布平台离线事件。"],
        ["comm/33007_V2", "33007_COMM_002", "心跳超时", "在线态长时间未收到控制域 0x83 心跳应答。", "关闭 TCP，发布离线事件并等待重连。"],
        ["comm/33007_V2", "33007_COMM_003", "平台帧解析失败", "启动符、长度、地址、COT 或 ASDU 尾部解析失败。", "丢弃该帧，必要时回 COT 0x06 拒绝。"],
        ["comm/33007_V2", "33007_COMM_004", "控制确认超时", "0x41 启停命令已下发但超时未收到 start_complete/stop_complete。", "回 COT 0x06，清理 PendingControl。"],
        ["comm/33007_V2", "33007_COMM_005", "计费模型缺失或解析失败", "0x50 未下发、费率字段非法，或远程启动时无可用 FeeModel。", "启动命令补充失败时记录日志；费率帧解析失败时回拒绝确认。"],
        ["comm/33007_V2", "33007_COMM_006", "交易记录确认失败", "0x42 确认帧无法匹配本地 pendingRecordTradeNo。", "保留未确认记录，等待补传或人工排查。"],
        ["comm/33007_V2", "33007_COMM_007", "平台发送阻塞/失败", "send 返回 EAGAIN/EWOULDBLOCK 或系统发送错误。", "记录发送异常，严重时关闭链路重连。"],
    ])


def section_deploy_test(doc):
    add_heading(doc, "四、配置、部署与测试要求", 1)
    add_heading(doc, "4.1 构建与部署", 2)
    add_bullets(doc, [
        "三个平台目录均提供 `Makefile.cross`，远端编译优先使用 `build imx6ul Makefile.cross`。",
        "目标机运行配置入口统一为 `/usr/app/config/tcu_comm.ini`，按部署平台替换对应配置内容。",
        "中石化平台配置可参考 `config/tcu_comm_zhongshihua2.ini`；远洋平台配置可参考 `core/comm/YuanyangV2/doc/yuanyang_v2_comm_config.example.ini`；33007_V2 配置可参考 `core/comm/33007_V2/doc/33007_v2_comm_config.example.ini`。",
        "平台通信进程需要内部 MQTT broker 可用，并与 `tcu_logic`、`pile_controller`、`meter` 的 topic 前缀和枪号偏置保持一致。",
    ])
    add_heading(doc, "4.2 联调测试项", 2)
    add_table(doc, ["测试项", "中石化2.0", "YuanyangV2", "33007_V2"], [
        ["登录/认证", "验证 0x01/0x02、0x0D/0x0A、0x0B/0x0C、0x03/0x04 状态机闭环。", "验证 0x81 公钥、0x82 认证、AES 业务加解密、心跳控制域闭环。", "验证 TCP 建链后直接在线，控制域 0x43/0x83 心跳闭环。"],
        ["平台启动", "验证 0xA8 或 0xA4 解析、费率保存、`start_charge` 发布、启动应答。", "验证 0x41 启动发布、PendingControl、`start_complete` 后 COT 0x04 确认。", "验证标准 0x41 启动控制解析，启动时附加 0x50 费率，完成后按原命令体回 COT 0x04。"],
        ["平台停机", "验证 0x36 解析、`stop_charge` 发布、0x35 应答。", "验证 0x41 停止发布、`stop_complete` 后确认。", "验证 0x41 停止控制发布，`stop_complete` 后 COT 0x04 确认。"],
        ["实时数据", "验证 0x13、0x23、0x25 等周期上报字段与 pile/meter/feeData 一致。", "验证总召和周期 0x27、0x25、0x10、0x11、0x12 上报顺序与点值。", "验证总召和周期 0x0B、0x27、0x05、0x25、0x10、0x11、0x12 明文上报顺序与点值。"],
        ["费率模型", "验证 0x0A、0x58 解析落库及每枪缓存更新。", "验证 0x51 主动请求、平台下发解析、保存 FeeModel、回确认。", "验证 0x50 主动请求、平台下发解析、保存 FeeModel、回确认，并在远程启动中携带费率。"],
        ["交易记录", "验证 update_record 触发 0x3D，上位 0x40 确认后本地确认。", "验证 update_record 触发 0x42，平台 COT 0x04 后发布 record_cfm 并确认本地记录。", "验证 update_record 触发 0x42，平台 COT 0x04 后发布 record_cfm 并确认本地记录。"],
        ["异常恢复", "断网、心跳超时、登录失败、密钥重置、解析失败均应有日志和重试/拒绝路径。", "断网、心跳超时、0x83 重置密钥、控制超时、费率解析失败均应闭环处理。", "断网、心跳超时、控制超时、费率缺失、发送阻塞、解析失败均应闭环处理。"],
    ], [1.0, 2.15, 2.15, 2.1])
    add_heading(doc, "4.3 后续扩展项", 2)
    add_bullets(doc, [
        "中石化 2.0：继续补齐协议附件中未接入帧，完善 SM2/SM4 强安全实现验收、幂等去重表和更多平台错误码映射。",
        "YuanyangV2：二期补充 0x2C/0x3C VIN/车牌请求启动、0x30/0x31/0x32 BMS、0x94/0x3D 刷卡启动；三期补充 0x71 档案、0x72 固件、0x33 有序充电、0xA1/0xA2 边缘网关数据。",
        "33007_V2：根据现场点表继续校准普通/扩展遥信遥测点号、故障码映射、功率调控字段和 0x50 费率模型边界条件。",
        "三个平台均应在联调后同步更新模块文档和测试用例，确保详细设计与实际代码保持一致。",
    ])


def main():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_revision(doc)
    add_toc(doc)
    section_intro(doc)
    section_overview(doc)
    section_zhong(doc)
    section_yuanyang(doc)
    section_33007(doc)
    section_deploy_test(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
